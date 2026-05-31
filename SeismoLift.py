from __future__ import annotations

import math
import os
import sys
from io import BytesIO
from dataclasses import dataclass
from datetime import datetime
from typing import Dict, List, Optional
from xml.sax.saxutils import escape

import pandas as pd
from unidecode import unidecode

import tkinter as tk
from tkinter import ttk, messagebox, filedialog

from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure

# Relatório .docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

PROGRAM_NAME = "SeismoLift"
PROGRAM_VERSION = "2025"
PROGRAM_URL = "https://github.com/lutondatomalela/SeismoLift-2025"
AUTHOR_NAME = "Engº Lutonda Tomalela"
APP_TITLE = PROGRAM_NAME


def resource_path(relative_path: str) -> str:
    """Resolve caminhos em execução normal e em executável PyInstaller onefile/onedir."""
    base_path = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)


def app_base_dir() -> str:
    """Directório da aplicação em execução normal ou congelada."""
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


BASE_DIR = app_base_dir()
ICON_DIR = resource_path("assets")
ICON_ICO = resource_path(os.path.join("assets", "seismolift_icon.ico"))
ICON_PNG = resource_path(os.path.join("assets", "seismolift_icon_32.png"))
ICON_PNG_LARGE = resource_path(os.path.join("assets", "seismolift_icon_128.png"))
DEFAULT_XLSX = resource_path("Zonas_Sismicas_PT.xlsx")
if not os.path.exists(DEFAULT_XLSX):
    DEFAULT_XLSX = os.path.join(BASE_DIR, "Zonas_Sismicas_PT.xlsx")
if not os.path.exists(DEFAULT_XLSX):
    DEFAULT_XLSX = os.path.join(BASE_DIR, "SeismoLift", "1_IN", "Zonas_Sismicas_PT.xlsx")

SHEETS = [
    "Portugal Continental",
    "Arquipélago da Madeira",
    "Arquipélago dos Açores",
]

# Classes de importância e coeficientes sísmicos segundo NP EN 1998-1:2010 - Anexo Nacional.
# Para Portugal Continental devem ser avaliadas as acções sísmicas Tipo 1 e Tipo 2,
# retendo-se o valor mais desfavorável para a categoria sísmica do ascensor.
CLASS_DESCRIPTIONS = {
    "I": "Importância menor",
    "II": "Edifícios correntes",
    "III": "Edifícios importantes",
    "IV": "Importância vital pós-sismo",
}

CLASS_OPTIONS = [f"{code} - {desc}" for code, desc in CLASS_DESCRIPTIONS.items()]

GAMMA_I_NA = {
    "Portugal Continental": {
        "Tipo 1": {"I": 0.65, "II": 1.00, "III": 1.45, "IV": 1.95},
        "Tipo 2": {"I": 0.75, "II": 1.00, "III": 1.25, "IV": 1.50},
    },
    "Arquipélago da Madeira": {
        "Tipo 1": {"I": 0.65, "II": 1.00, "III": 1.45, "IV": 1.95},
    },
    "Arquipélago dos Açores": {
        "Tipo 2": {"I": 0.85, "II": 1.00, "III": 1.15, "IV": 1.35},
    },
}

GAMMA_A_OPTIONS = {
    "Ascensor corrente (γa = 1,00)": 1.00,
    "Ascensor vital/bombeiros (γa = 1,50)": 1.50,
}

CALC_MODE_GENERAL = "Geral EC8 / EN 81-77"
CALC_MODE_ET11 = "ET 11/2020 - edifício de base fixa"
CALCULATION_MODES = [CALC_MODE_GENERAL, CALC_MODE_ET11]

STRUCTURE_OPTIONS_ELEVATOR = {
    "Pórticos metálicos": {"label": "Pórticos metálicos", "Ct": 0.085},
    "Pórticos de betão ou contraventados": {"label": "Pórticos de betão ou contraventados", "Ct": 0.075},
    "Estruturas em geral": {"label": "Estruturas em geral", "Ct": 0.050},
}


SOIL_DESCRIPTIONS = {
    "A": "Rocha/formação rochosa",
    "B": "Depósitos muito compactos/rijos",
    "C": "Depósitos densos/médios ou rijos",
    "D": "Solos soltos/médios ou coesivos moles/firmes",
    "E": "Aluvião superficial sobre substrato rígido",
}

SOIL_OPTIONS = [f"{code} - {desc}" for code, desc in SOIL_DESCRIPTIONS.items()]

# Parâmetros para a aba "Espetros de Resposta"
SOIL_PARAMS = {
    "Tipo 1": {
        "A": {"S": 1.00, "TB": 0.10, "TC": 0.60, "TD": 2.00},
        "B": {"S": 1.35, "TB": 0.10, "TC": 0.60, "TD": 2.00},
        "C": {"S": 1.60, "TB": 0.10, "TC": 0.60, "TD": 2.00},
        "D": {"S": 2.00, "TB": 0.10, "TC": 0.80, "TD": 2.00},
        "E": {"S": 1.80, "TB": 0.10, "TC": 0.60, "TD": 2.00},
    },
    "Tipo 2": {
        "A": {"S": 1.00, "TB": 0.10, "TC": 0.25, "TD": 2.00},
        "B": {"S": 1.35, "TB": 0.10, "TC": 0.25, "TD": 2.00},
        "C": {"S": 1.60, "TB": 0.10, "TC": 0.25, "TD": 2.00},
        "D": {"S": 2.00, "TB": 0.10, "TC": 0.30, "TD": 2.00},
        "E": {"S": 1.80, "TB": 0.10, "TC": 0.25, "TD": 2.00},
    },
}

STRUCTURE_OPTIONS_SPECTRUM = {
    "Pórticos metálicos": 0.085,
    "Pórticos de betão / contraventados": 0.075,
    "Estruturas em geral": 0.050,
}

T1_MODE_AUTO = "Automático por Ct·H^0,75"
T1_MODE_MANUAL = "Manual"
T1_MODE_OPTIONS = [T1_MODE_AUTO, T1_MODE_MANUAL]
SPECTRUM_UNIT_OPTIONS = ["m/s²", "g"]

QA_DEFAULT = 2.0
G = 9.81


@dataclass
class ZoneResult:
    input_name: str
    concelho: str
    regiao: str
    z1: Optional[float]
    z2: Optional[float]
    agR1: Optional[float]
    agR2: Optional[float]
    spectro_type: str
    agR_used: float



@dataclass
class ElevatorScenario:
    spectro_type: str
    z_sismica: Optional[float]
    agR: float
    gamma_l: float
    ag: float
    alfa: float
    S: float
    Sa: float
    ad: float
    categoria: int

@dataclass
class ElevatorResult:
    zone: ZoneResult
    calc_mode: str
    classe: str
    gamma_l: float
    gamma_a: float
    terreno: str
    S: float
    estrutura_codigo: str
    estrutura_label: str
    Ct: float
    H: float
    z: float
    T1: float
    Ta: float
    ag: float
    agR: float
    alfa: float
    qa: float
    Sa: float
    ad: float
    categoria: int
    cenarios: List[ElevatorScenario]
    gamma_a_note: str


@dataclass
class SpectrumCurve:
    spectro_type: str
    z_sismica: Optional[float]
    agR: float
    gamma_I: float
    ag: float
    alpha: float
    S: float
    TB: float
    TC: float
    TD: float
    periods: List[float]
    Se: List[float]
    Sde: List[float]


@dataclass
class SpectrumResult:
    zone: ZoneResult
    class_importance: str
    terrain_type: str
    structure_type: str
    gamma_I: float
    gamma_a: float
    ag: float
    alpha: float
    S: float
    TB: float
    TC: float
    TD: float
    Ct: float
    H: float
    z: float
    T1: float
    Ta: float
    Sa_ns: float
    ad: float
    categoria: int
    cenarios: List[ElevatorScenario]
    spectra: List[SpectrumCurve]
    gamma_a_note: str
    xi: float
    eta: float
    periods: List[float]
    Se: List[float]
    Sde: List[float]
    t1_mode: str


class SeismoLiftCore:
    def __init__(self, excel_path: str):
        self.excel_path = excel_path
        self.sheets: Dict[str, pd.DataFrame] = {}
        self._load()

    def _load(self) -> None:
        if not os.path.exists(self.excel_path):
            raise FileNotFoundError(f"Base de dados não encontrada: {self.excel_path}")
        self.sheets = {name: pd.read_excel(self.excel_path, sheet_name=name, header=0) for name in SHEETS}

    @staticmethod
    def _norm(text: str) -> str:
        return unidecode(str(text)).lower().strip()

    def find_zone(self, localidade: str) -> ZoneResult:
        localidade_n = self._norm(localidade)

        for regiao, df in self.sheets.items():
            col_b = df.iloc[1:, 1]
            for idx, concelho in col_b.items():
                if pd.isna(concelho):
                    continue
                if localidade_n == self._norm(concelho):
                    if regiao == "Portugal Continental":
                        vals = df.iloc[idx, 2:6].tolist()
                        z1, agR1, z2, agR2 = vals[0], float(vals[1]), vals[2], float(vals[3])
                        spectro_type = "Tipo 1"
                        agR_used = agR1
                    elif regiao == "Arquipélago da Madeira":
                        vals = df.iloc[idx, 2:4].tolist()
                        z1, agR1, z2, agR2 = float(vals[0]), float(vals[1]), None, None
                        spectro_type = "Tipo 1"
                        agR_used = agR1
                    else:
                        vals = df.iloc[idx, 2:4].tolist()
                        z1, agR1, z2, agR2 = None, None, float(vals[0]), float(vals[1])
                        spectro_type = "Tipo 2"
                        agR_used = agR2

                    return ZoneResult(
                        input_name=localidade,
                        concelho=str(concelho),
                        regiao=regiao,
                        z1=z1,
                        z2=z2,
                        agR1=agR1,
                        agR2=agR2,
                        spectro_type=spectro_type,
                        agR_used=agR_used,
                    )

        raise ValueError(f"Concelho não encontrado: {localidade}")

    @staticmethod
    def soil_coefficient_original(terreno: str) -> float:
        return {"A": 1.00, "B": 1.35, "C": 1.60, "D": 2.00, "E": 1.80}.get(terreno, 1.60)

    @staticmethod
    def category_from_ad(ad: float) -> int:
        if ad <= 1.0:
            return 0
        if ad <= 2.5:
            return 1
        if ad <= 4.0:
            return 2
        return 3

    @staticmethod
    def category_note(categoria: int) -> str:
        return {
            0: "Não são requeridas acções sísmicas adicionais da EN 81-77.",
            1: "São requeridas acções correctivas de baixa expressão.",
            2: "São requeridas acções correctivas de média expressão.",
            3: "São requeridas acções correctivas importantes.",
        }.get(categoria, "Categoria sísmica não definida.")

    @staticmethod
    def checklist_for_category(categoria: int) -> List[str]:
        base = [
            "Confirmar que o cálculo de ad foi fornecido/validado pelo projectista da estrutura ou dono de obra.",
            "Confirmar o cumprimento das normas base aplicáveis ao ascensor, nomeadamente EN 81-20 e EN 81-50.",
        ]
        if categoria == 0:
            return base + ["Categoria 0: a EN 81-77 não introduz medidas sísmicas adicionais específicas."]
        items = base + [
            "Avaliar riscos de engate de cabos, correntes ou meios de compensação em elementos salientes na caixa.",
            "Verificar dispositivos de retenção da cabina e do contrapeso/balanço, quando aplicável.",
            "Confirmar protecção contra saída de cabos das rodas de tracção, desvio, polias e rodas dentadas.",
            "Confirmar retenção dos blocos do contrapeso e adequação das fixações.",
            "Verificar fixação de maquinaria, quadros de comando e equipamento eléctrico instalado na caixa.",
        ]
        if categoria >= 2:
            items += [
                "Verificar tensões e flechas admissíveis nas guias durante a acção sísmica.",
                "Avaliar comportamento do ascensor em modo sísmico e sequência de paragem/rearranque.",
                "Confirmar requisitos funcionais associados a falha de alimentação principal.",
            ]
        if categoria >= 3:
            items += [
                "Prever sistema de detecção sísmica quando exigido pela EN 81-77 e pela especificação do projecto.",
                "Definir posição e condições de instalação do sistema de detecção sísmica em articulação com o edifício.",
                "Confirmar indicação visual do modo sísmico e procedimentos de reposição ao serviço.",
            ]
        return items

    @staticmethod
    def gamma_a_from_option(option: str) -> float:
        return GAMMA_A_OPTIONS.get(option, 1.00)

    @staticmethod
    def gamma_i(regiao: str, spectro_type: str, classe: str) -> float:
        return GAMMA_I_NA[regiao][spectro_type][classe]

    @staticmethod
    def terrain_code(terreno: str) -> str:
        value = str(terreno).strip()
        if value in SOIL_DESCRIPTIONS:
            return value
        if value:
            code = value[0].upper()
            if code in SOIL_DESCRIPTIONS:
                return code
        raise ValueError("Tipo de terreno inválido.")

    @staticmethod
    def class_code(classe: str) -> str:
        value = str(classe).strip()
        if value in CLASS_DESCRIPTIONS:
            return value
        if value:
            code = value.split("-", 1)[0].strip().upper()
            if code in CLASS_DESCRIPTIONS:
                return code
        raise ValueError("Classe de importância inválida.")

    @staticmethod
    def class_label(classe: str) -> str:
        code = SeismoLiftCore.class_code(classe)
        return f"{code} - {CLASS_DESCRIPTIONS[code]}"

    @staticmethod
    def terrain_label(terreno: str) -> str:
        code = SeismoLiftCore.terrain_code(terreno)
        return f"{code} - {SOIL_DESCRIPTIONS[code]}"

    def available_seismic_actions(self, zone: ZoneResult) -> List[tuple[str, Optional[float], float]]:
        actions: List[tuple[str, Optional[float], float]] = []
        if zone.agR1 is not None:
            actions.append(("Tipo 1", zone.z1, float(zone.agR1)))
        if zone.agR2 is not None:
            actions.append(("Tipo 2", zone.z2, float(zone.agR2)))
        return actions

    def compute_nonstructural_scenario(self, regiao: str, spectro_type: str, z_sismica: Optional[float], agR: float,
                                       classe: str, terreno: str, gamma_a: float, qa: float,
                                       H: float, z: float, T1: float,
                                       calc_mode: str = CALC_MODE_GENERAL, Ta: float = 0.0) -> ElevatorScenario:
        gamma_l = self.gamma_i(regiao, spectro_type, classe)
        ag = gamma_l * agR
        alfa = ag / G
        S = self.soil_coefficient_original(terreno)

        if calc_mode == CALC_MODE_ET11:
            # ET 11/2020, edifício de base fixa: ad = 1,25·γI·agR·S para γa = 1,00.
            # Mantém-se γa explícito para os casos de ascensor vital/bombeiros.
            Sa = 2.5 * alfa * S
        else:
            if T1 <= 0:
                raise ValueError("T1 deve ser positivo no modo geral EC8 / EN 81-77.")
            Sa = alfa * S * (3 * (1 + z / H) / (1 + (1 - Ta / T1) ** 2) - 0.5)
            Sa = max(Sa, alfa * S)

        ad = Sa * (gamma_a / qa) * G
        return ElevatorScenario(
            spectro_type=spectro_type,
            z_sismica=z_sismica,
            agR=agR,
            gamma_l=gamma_l,
            ag=ag,
            alfa=alfa,
            S=S,
            Sa=Sa,
            ad=ad,
            categoria=self.category_from_ad(ad),
        )

    def compute_elevator(self, localidade: str, classe: str, terreno: str, estrutura_codigo: str,
                         H: float, z: float, gamma_a_option: str = "Ascensor corrente (γa = 1,00)",
                         calc_mode: str = CALC_MODE_GENERAL) -> ElevatorResult:
        zone = self.find_zone(localidade)
        if H <= 0:
            raise ValueError("H deve ser positivo.")
        if z <= 0:
            raise ValueError("z deve ser positivo.")
        if z > H:
            raise ValueError("z não pode ser superior a H.")
        if calc_mode not in CALCULATION_MODES:
            raise ValueError("Modo de cálculo inválido.")
        classe = self.class_code(classe)
        terreno = self.terrain_code(terreno)
        if estrutura_codigo not in STRUCTURE_OPTIONS_ELEVATOR:
            raise ValueError("Tipo de estrutura inválido.")

        gamma_a = self.gamma_a_from_option(gamma_a_option)
        estrutura_label = STRUCTURE_OPTIONS_ELEVATOR[estrutura_codigo]["label"]
        Ct = STRUCTURE_OPTIONS_ELEVATOR[estrutura_codigo]["Ct"]
        qa = QA_DEFAULT
        T1 = Ct * H ** (3 / 4)
        Ta = 0.0

        cenarios = [
            self.compute_nonstructural_scenario(zone.regiao, st, zs, agr, classe, terreno, gamma_a, qa, H, z, T1, calc_mode)
            for st, zs, agr in self.available_seismic_actions(zone)
        ]
        if not cenarios:
            raise ValueError("Não foram encontradas acções sísmicas válidas para a localidade.")

        # Critério ET 11/2020 / NP EN 1998-1: para o Continente avaliam-se Tipo 1 e Tipo 2 e adopta-se o valor mais desfavorável.
        controlling = max(cenarios, key=lambda c: c.ad)
        zone.spectro_type = controlling.spectro_type
        zone.agR_used = controlling.agR

        return ElevatorResult(
            zone=zone,
            calc_mode=calc_mode,
            classe=classe,
            gamma_l=controlling.gamma_l,
            gamma_a=gamma_a,
            terreno=terreno,
            S=controlling.S,
            estrutura_codigo=estrutura_codigo,
            estrutura_label=estrutura_label,
            Ct=Ct,
            H=H,
            z=z,
            T1=T1,
            Ta=Ta,
            ag=controlling.ag,
            agR=controlling.agR,
            alfa=controlling.alfa,
            qa=qa,
            Sa=controlling.Sa,
            ad=controlling.ad,
            categoria=controlling.categoria,
            cenarios=cenarios,
            gamma_a_note=gamma_a_option,
        )

    def compute_spectrum(self, localidade: str, classe: str, terreno: str, estrutura: str,
                         H: float, z: float, xi: float = 5.0,
                         t_max: float = 4.0, n_points: int = 401,
                         gamma_a_option: str = "Ascensor corrente (γa = 1,00)",
                         Ta: float = 0.0, t1_mode: str = T1_MODE_AUTO,
                         T1_manual: Optional[float] = None) -> SpectrumResult:
        zone = self.find_zone(localidade)

        if H <= 0:
            raise ValueError("H deve ser positivo.")
        if z < 0:
            raise ValueError("z não pode ser negativo.")
        if z > H:
            raise ValueError("z não pode ser superior a H.")
        if Ta < 0:
            raise ValueError("Ta não pode ser negativo.")
        if xi <= 0:
            raise ValueError("O amortecimento deve ser positivo.")
        if t_max <= 0:
            raise ValueError("T máx deve ser positivo.")
        if n_points < 10:
            raise ValueError("O número de pontos do espetro deve ser pelo menos 10.")
        terreno = self.terrain_code(terreno)
        if estrutura not in STRUCTURE_OPTIONS_SPECTRUM:
            raise ValueError(f"Tipo de estrutura inválido: {estrutura}")
        classe = self.class_code(classe)

        gamma_a = self.gamma_a_from_option(gamma_a_option)
        Ct = STRUCTURE_OPTIONS_SPECTRUM[estrutura]
        if t1_mode == T1_MODE_MANUAL:
            if T1_manual is None:
                raise ValueError("Introduza T1 manual.")
            T1 = float(T1_manual)
        else:
            t1_mode = T1_MODE_AUTO
            T1 = Ct * (H ** 0.75)
        if T1 <= 0:
            raise ValueError("T1 deve ser positivo.")

        periods = [i * t_max / (n_points - 1) for i in range(n_points)]
        eta = max(math.sqrt(10.0 / (5.0 + xi)), 0.55)

        spectra: List[SpectrumCurve] = []
        cenarios: List[ElevatorScenario] = []
        for st, zs, agr in self.available_seismic_actions(zone):
            scenario = self.compute_nonstructural_scenario(
                zone.regiao, st, zs, agr, classe, terreno, gamma_a, QA_DEFAULT, H, z, T1,
                CALC_MODE_GENERAL, Ta=Ta,
            )
            cenarios.append(scenario)
            p = SOIL_PARAMS[st][terreno]
            S, TB, TC, TD = p["S"], p["TB"], p["TC"], p["TD"]
            curve = SpectrumCurve(
                spectro_type=st,
                z_sismica=zs,
                agR=agr,
                gamma_I=scenario.gamma_l,
                ag=scenario.ag,
                alpha=scenario.alfa,
                S=S,
                TB=TB,
                TC=TC,
                TD=TD,
                periods=periods,
                Se=[self.elastic_spectrum(t, scenario.ag, S, TB, TC, TD, eta) for t in periods],
                Sde=[],
            )
            curve.Sde = [0.0 if t == 0 else se * (t / (2 * math.pi)) ** 2 for t, se in zip(curve.periods, curve.Se)]
            spectra.append(curve)

        if not cenarios:
            raise ValueError("Não foram encontradas acções sísmicas válidas para a localidade.")
        controlling = max(cenarios, key=lambda c: c.ad)
        zone.spectro_type = controlling.spectro_type
        zone.agR_used = controlling.agR

        controlling_curve = next(c for c in spectra if c.spectro_type == controlling.spectro_type)

        return SpectrumResult(
            zone=zone,
            class_importance=classe,
            terrain_type=terreno,
            structure_type=estrutura,
            gamma_I=controlling.gamma_l,
            gamma_a=gamma_a,
            ag=controlling.ag,
            alpha=controlling.alfa,
            S=controlling_curve.S,
            TB=controlling_curve.TB,
            TC=controlling_curve.TC,
            TD=controlling_curve.TD,
            Ct=Ct,
            H=H,
            z=z,
            T1=T1,
            Ta=Ta,
            Sa_ns=controlling.Sa,
            ad=controlling.ad,
            categoria=controlling.categoria,
            cenarios=cenarios,
            spectra=spectra,
            gamma_a_note=gamma_a_option,
            xi=xi,
            eta=eta,
            periods=periods,
            Se=controlling_curve.Se,
            Sde=controlling_curve.Sde,
            t1_mode=t1_mode,
        )

    @staticmethod
    def elastic_spectrum(T: float, ag: float, S: float, TB: float, TC: float, TD: float, eta: float) -> float:
        if T <= TB:
            return ag * S * (1.0 + (T / TB) * (2.5 * eta - 1.0))
        if T <= TC:
            return ag * S * 2.5 * eta
        if T <= TD:
            return ag * S * 2.5 * eta * (TC / T)
        return ag * S * 2.5 * eta * (TC * TD / (T ** 2))

    @staticmethod
    def spectral_displacement(Se: float, T: float) -> float:
        if T <= 0:
            return 0.0
        return Se * (T / (2 * math.pi)) ** 2

    def curve_values_at(self, curve: SpectrumCurve, T: float, eta: Optional[float] = None) -> tuple[float, float]:
        if eta is None:
            eta = 1.0
        Se = self.elastic_spectrum(T, curve.ag, curve.S, curve.TB, curve.TC, curve.TD, eta)
        return Se, self.spectral_displacement(Se, T)


class BaseTab(ttk.Frame):
    def __init__(self, parent, app: "SeismoLiftGUI"):
        super().__init__(parent)
        self.app = app
        self.core = app.core

    @staticmethod
    def _fmt(value: object, digits: int = 3) -> str:
        if value is None:
            return "-"
        if isinstance(value, float):
            return f"{value:.{digits}f}".replace(".", ",")
        return str(value)

    @staticmethod
    def _set_run_font(run, font_name: str = "Courier New", size: Optional[int] = None, bold: Optional[bool] = None) -> None:
        run.font.name = font_name
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(attr), font_name)

    @staticmethod
    def _add_hidden_hyperlink(paragraph, text: str, url: str, font_name: str = "Courier New", size: int = 8, bold: bool = False):
        """Add a clickable hyperlink while displaying only the provided text.

        The hyperlink is intentionally formatted like normal text (black, no underline),
        so the URL remains hidden behind the program name in the report.
        """
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")

        r_fonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(attr), font_name)
        r_pr.append(r_fonts)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size * 2))
        r_pr.append(sz)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        r_pr.append(color)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "none")
        r_pr.append(underline)

        if bold:
            r_pr.append(OxmlElement("w:b"))

        new_run.append(r_pr)
        text_element = OxmlElement("w:t")
        text_element.text = text
        new_run.append(text_element)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    @staticmethod
    def parse_float(value: str, name: str) -> float:
        try:
            return float(value.replace(",", ".").strip())
        except Exception as exc:
            raise ValueError(f"Valor inválido para {name}: {value}") from exc


class ScrollableFrame(ttk.Frame):
    """Frame com scroll vertical para manter a GUI usável em ecrãs menores."""

    def __init__(self, parent, width: int = 390, *args, **kwargs):
        super().__init__(parent, *args, **kwargs)
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)

        self.canvas = tk.Canvas(self, highlightthickness=0, borderwidth=0, width=width)
        self.vscroll = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = ttk.Frame(self.canvas, padding=(0, 0, 8, 0))
        self.window_id = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")

        self.canvas.configure(yscrollcommand=self.vscroll.set)
        self.canvas.grid(row=0, column=0, sticky="nsew")
        self.vscroll.grid(row=0, column=1, sticky="ns")

        self.scrollable_frame.bind("<Configure>", self._on_frame_configure)
        self.canvas.bind("<Configure>", self._on_canvas_configure)
        self.canvas.bind("<Enter>", self._bind_mousewheel)
        self.canvas.bind("<Leave>", self._unbind_mousewheel)

    def _on_frame_configure(self, _event=None):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def _on_canvas_configure(self, event):
        self.canvas.itemconfigure(self.window_id, width=max(event.width - 2, 270))

    def _bind_mousewheel(self, _event=None):
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        self.canvas.bind_all("<Button-4>", self._on_mousewheel_linux)
        self.canvas.bind_all("<Button-5>", self._on_mousewheel_linux)

    def _unbind_mousewheel(self, _event=None):
        self.canvas.unbind_all("<MouseWheel>")
        self.canvas.unbind_all("<Button-4>")
        self.canvas.unbind_all("<Button-5>")

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def _on_mousewheel_linux(self, event):
        self.canvas.yview_scroll(-1 if event.num == 4 else 1, "units")


class AutoPanedWindow(ttk.Panedwindow):
    """PanedWindow com posição inicial de sash, sem impedir o redimensionamento pelo utilizador."""

    def __init__(self, parent, orient="horizontal", sash_position: Optional[int] = None, *args, **kwargs):
        super().__init__(parent, orient=orient, *args, **kwargs)
        self._initial_sash_position = sash_position
        if sash_position is not None:
            self.after_idle(self._apply_initial_sash_position)

    def _apply_initial_sash_position(self):
        try:
            if self.panes():
                self.sashpos(0, self._initial_sash_position)
        except tk.TclError:
            pass


class ElevatorTab(BaseTab):
    def __init__(self, parent, app: "SeismoLiftGUI"):
        super().__init__(parent, app)
        self.result: Optional[ElevatorResult] = None
        self._build()

    def _build(self):
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)

        paned = AutoPanedWindow(self, orient="horizontal", sash_position=410)
        paned.grid(row=0, column=0, sticky="nsew")

        left_scroll = ScrollableFrame(paned, width=390)
        left = left_scroll.scrollable_frame
        left.columnconfigure(0, weight=1)

        right = ttk.Frame(paned, padding=(10, 12, 12, 12))
        right.columnconfigure(0, weight=1)
        right.rowconfigure(1, weight=1)

        paned.add(left_scroll, weight=0)
        paned.add(right, weight=1)

        ttk.Label(left, text="Categoria Sísmica de Elevadores", style="Title.TLabel").grid(row=0, column=0, sticky="w")

        form = ttk.LabelFrame(left, text="Entradas", padding=10)
        form.grid(row=1, column=0, sticky="nwe", pady=(12, 0))
        form.columnconfigure(1, weight=1)

        self.localidade_var = tk.StringVar()
        self.classe_var = tk.StringVar(value=self.core.class_label("II"))
        self.terreno_var = tk.StringVar(value=self.core.terrain_label("B"))
        self.estrutura_var = tk.StringVar(value="Estruturas em geral")
        self.gamma_a_var = tk.StringVar(value="Ascensor corrente (γa = 1,00)")
        self.calc_mode_var = tk.StringVar(value=CALC_MODE_GENERAL)
        self.H_var = tk.StringVar(value="12,0")
        self.z_var = tk.StringVar(value="12,0")

        labels = ["Concelho", "Classe", "Terreno", "Estrutura", "γa", "Modo de cálculo", "H [m]", "z [m]"]
        for i, txt in enumerate(labels):
            ttk.Label(form, text=txt).grid(row=i, column=0, sticky="w", padx=(0, 8), pady=4)

        ttk.Entry(form, textvariable=self.localidade_var).grid(row=0, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Combobox(form, textvariable=self.classe_var, state="readonly", values=CLASS_OPTIONS, width=34).grid(row=1, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Combobox(form, textvariable=self.terreno_var, state="readonly", values=SOIL_OPTIONS, width=34).grid(row=2, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Combobox(form, textvariable=self.estrutura_var, state="readonly", values=list(STRUCTURE_OPTIONS_ELEVATOR.keys()), width=34).grid(row=3, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Combobox(form, textvariable=self.gamma_a_var, state="readonly", values=list(GAMMA_A_OPTIONS.keys()), width=34).grid(row=4, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Combobox(form, textvariable=self.calc_mode_var, state="readonly", values=CALCULATION_MODES, width=34).grid(row=5, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Entry(form, textvariable=self.H_var).grid(row=6, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Entry(form, textvariable=self.z_var).grid(row=7, column=1, columnspan=2, sticky="ew", pady=4)

        buttons = ttk.Frame(left)
        buttons.grid(row=3, column=0, sticky="ew", pady=(12, 0))
        buttons.columnconfigure((0, 1), weight=1)
        ttk.Button(buttons, text="Calcular", command=self.calculate).grid(row=0, column=0, sticky="ew", padx=(0, 6))
        ttk.Button(buttons, text="Limpar", command=self.clear_all).grid(row=0, column=1, sticky="ew")
        ttk.Button(buttons, text="Exportar relatório .docx", command=lambda: self.export_report("docx")).grid(row=1, column=0, sticky="ew", pady=(8, 0), padx=(0, 6))
        ttk.Button(buttons, text="Exportar relatório .pdf", command=lambda: self.export_report("pdf")).grid(row=1, column=1, sticky="ew", pady=(8, 0))
        ttk.Button(buttons, text="Exportar dados .xlsx", command=self.export_xlsx).grid(row=2, column=0, columnspan=2, sticky="ew", pady=(8, 0))

        self.category_badge = tk.Label(left, text="Categoria sísmica: —", font=("Segoe UI", 11, "bold"), bg="#f0f0f0", fg="#000", padx=10, pady=10, relief="groove")
        self.category_badge.grid(row=4, column=0, sticky="ew", pady=(12, 0))

        notes = ttk.LabelFrame(left, text="Notas", padding=10)
        notes.grid(row=5, column=0, sticky="new", pady=(12, 0))
        ttk.Label(
            notes,
            text=(
                "• Portugal Continental: avalia Tipo 1 e Tipo 2 e adopta o valor mais desfavorável.\n"
                "• A categoria sísmica do elevador usa qa = 2,0.\n"
                "• γa = 1,0 por defeito; γa = 1,5 para ascensor vital/bombeiros."
            ),
            justify="left", wraplength=340,
        ).grid(row=0, column=0, sticky="w")

        header = ttk.Frame(right)
        header.grid(row=0, column=0, sticky="ew", pady=(0, 8))
        header.columnconfigure(0, weight=1)
        ttk.Label(header, text="Resultados", style="Title.TLabel").grid(row=0, column=0, sticky="w")

        self.tree = ttk.Treeview(right, columns=("valor",), show="tree headings", style="Result.Treeview", height=20)
        self.tree.heading("#0", text="Parâmetro")
        self.tree.heading("valor", text="Valor")
        self.tree.column("#0", width=280, minwidth=180, anchor="w", stretch=False)
        self.tree.column("valor", width=520, minwidth=220, anchor="w", stretch=True)
        self.tree.grid(row=1, column=0, sticky="nsew")
        yscroll = ttk.Scrollbar(right, orient="vertical", command=self.tree.yview)
        xscroll = ttk.Scrollbar(right, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=yscroll.set, xscrollcommand=xscroll.set)
        yscroll.grid(row=1, column=1, sticky="ns")
        xscroll.grid(row=2, column=0, sticky="ew")

    def calculate(self):
        try:
            result = self.core.compute_elevator(
                localidade=self.localidade_var.get().strip(),
                classe=self.classe_var.get().strip(),
                terreno=self.terreno_var.get().strip(),
                estrutura_codigo=self.estrutura_var.get(),
                H=self.parse_float(self.H_var.get(), "H"),
                z=self.parse_float(self.z_var.get(), "z"),
                gamma_a_option=self.gamma_a_var.get(),
                calc_mode=self.calc_mode_var.get(),
            )
        except Exception as exc:
            messagebox.showerror("Erro", str(exc))
            return

        self.result = result
        self.populate_results()

    def populate_results(self):
        assert self.result is not None
        r = self.result
        for item in self.tree.get_children():
            self.tree.delete(item)

        data = [
            ("Localização", f"{r.zone.concelho}, {r.zone.regiao}"),
            ("Acção sísmica condicionante", r.zone.spectro_type),
            ("Modo de cálculo", r.calc_mode),
            ("Classe de importância", self.core.class_label(r.classe)),
            ("Coeficiente de importância γI [-]", self._fmt(r.gamma_l)),
            ("Aceleração de referência agR [m/s²]", self._fmt(r.agR)),
            ("Aceleração de cálculo ag [m/s²]", self._fmt(r.ag)),
            ("Tipo de terreno", self.core.terrain_label(r.terreno)),
            ("Coeficiente de solo S [-]", self._fmt(r.S)),
            ("Tipo de estrutura", r.estrutura_label),
            ("Coeficiente Ct [-]", self._fmt(r.Ct)),
            ("Altura total do edifício H [m]", self._fmt(r.H)),
            ("Altura do ponto considerado z [m]", self._fmt(r.z)),
            ("Relação α = ag/g [-]", self._fmt(r.alfa)),
            ("Período fundamental do edifício T1 [s]", self._fmt(r.T1)),
            ("Coeficiente de comportamento qa [-]", self._fmt(r.qa)),
            ("Coeficiente γa [-]", self._fmt(r.gamma_a)),
            ("Critério γa", r.gamma_a_note),
            ("Ordenada espectral Sa [-]", self._fmt(r.Sa)),
            ("Aceleração de projecto ad [m/s²]", self._fmt(r.ad)),
            ("Categoria sísmica do elevador", str(r.categoria)),
        ]

        for k, v in data:
            self.tree.insert("", "end", text=k, values=(v,))
        if len(r.cenarios) > 1:
            self.tree.insert("", "end", text="--- Acções avaliadas ---", values=("",))
            for c in r.cenarios:
                self.tree.insert("", "end", text=f"{c.spectro_type}: ad [m/s²] / categoria", values=(f"{self._fmt(c.ad)} / {c.categoria}",))

        self._update_category_badge(r.categoria)

    def clear_all(self):
        self.localidade_var.set("")
        self.classe_var.set(self.core.class_label("II"))
        self.terreno_var.set(self.core.terrain_label("B"))
        self.estrutura_var.set("Estruturas em geral")
        self.gamma_a_var.set("Ascensor corrente (γa = 1,00)")
        self.calc_mode_var.set(CALC_MODE_GENERAL)
        self.H_var.set("12,0")
        self.z_var.set("12,0")
        self.result = None
        for item in self.tree.get_children():
            self.tree.delete(item)
        self._update_category_badge(None)

    def _update_category_badge(self, categoria: Optional[int]):
        colors_map = {
            None: ("Categoria sísmica: —", "#f0f0f0", "#000000"),
            0: ("Categoria sísmica: 0", "#90EE90", "#000000"),
            1: ("Categoria sísmica: 1", "#FFFF00", "#000000"),
            2: ("Categoria sísmica: 2", "#FFA500", "#000000"),
            3: ("Categoria sísmica: 3", "#FF6B6B", "#000000"),
        }
        txt, bg, fg = colors_map.get(categoria, colors_map[None])
        self.category_badge.config(text=txt, bg=bg, fg=fg)

    def export_report(self, fmt: str = "docx"):
        if not self.result:
            messagebox.showwarning("Sem dados", "Calcule primeiro a categoria sísmica.")
            return
        try:
            ext = ".docx" if fmt == "docx" else ".pdf"
            default_name = f"SeismoLift_report_{self.result.zone.concelho.replace(' ', '_')}{ext}"
            filetypes = [("Documento Word", "*.docx")] if fmt == "docx" else [("Documento PDF", "*.pdf")]
            path = filedialog.asksaveasfilename(
                title="Guardar relatório",
                defaultextension=ext,
                initialfile=default_name,
                filetypes=filetypes,
            )
            if not path:
                return
            if fmt == "docx":
                self._generate_report_docx(path)
            else:
                self._generate_report_pdf(path)
            messagebox.showinfo("Sucesso", f"Relatório guardado em:\n{path}")
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível exportar o relatório.\n\n{exc}")

    def export_xlsx(self):
        if not self.result:
            messagebox.showwarning("Sem dados", "Calcule primeiro a categoria sísmica.")
            return
        try:
            default_name = f"SeismoLift_dados_{self.result.zone.concelho.replace(' ', '_')}.xlsx"
            path = filedialog.asksaveasfilename(
                title="Guardar dados de cálculo",
                defaultextension=".xlsx",
                initialfile=default_name,
                filetypes=[("Livro Excel", "*.xlsx")],
            )
            if not path:
                return
            self._generate_report_xlsx(path)
            messagebox.showinfo("Sucesso", f"Dados guardados em:\n{path}")
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível exportar os dados.\n\n{exc}")

    def _generate_report_xlsx(self, path: str):
        assert self.result is not None
        r = self.result
        checklist = self.core.checklist_for_category(r.categoria)
        summary = [
            ["Programa", PROGRAM_NAME],
            ["Local", f"{r.zone.concelho}, {r.zone.regiao}"],
            ["Modo de cálculo", r.calc_mode],
            ["Classe de importância", self.core.class_label(r.classe)],
            ["Tipo de terreno", self.core.terrain_label(r.terreno)],
            ["Acção sísmica condicionante", r.zone.spectro_type],
            ["ad [m/s²]", r.ad],
            ["Categoria sísmica", r.categoria],
        ]
        inputs = [
            ["γI [-]", r.gamma_l], ["γa [-]", r.gamma_a], ["qa [-]", r.qa],
            ["agR [m/s²]", r.agR], ["ag [m/s²]", r.ag], ["α = ag/g [-]", r.alfa],
            ["S [-]", r.S], ["Ct [-]", r.Ct], ["H [m]", r.H], ["z [m]", r.z],
            ["T1 [s]", r.T1], ["Ta [s]", r.Ta], ["Sa [-]", r.Sa],
        ]
        scenarios = [["Acção", "Zona", "γI", "agR [m/s²]", "ag [m/s²]", "S", "Sa", "ad [m/s²]", "Categoria"]]
        for c in r.cenarios:
            scenarios.append([c.spectro_type, c.z_sismica, c.gamma_l, c.agR, c.ag, c.S, c.Sa, c.ad, c.categoria])
        with pd.ExcelWriter(path, engine="openpyxl") as writer:
            pd.DataFrame(summary, columns=["Campo", "Valor"]).to_excel(writer, index=False, sheet_name="Resumo")
            pd.DataFrame(inputs, columns=["Parâmetro", "Valor"]).to_excel(writer, index=False, sheet_name="Parâmetros")
            pd.DataFrame(scenarios[1:], columns=scenarios[0]).to_excel(writer, index=False, sheet_name="Acções")
            pd.DataFrame([[i + 1, item] for i, item in enumerate(checklist)], columns=["N.º", "Requisito/controlo"]).to_excel(writer, index=False, sheet_name="Checklist EN 81-77")

            # Metadados do livro Excel.
            props = writer.book.properties
            props.title = "SeismoLift - Dados de cálculo"
            props.subject = "Cálculo da aceleração de projecto e categoria sísmica do ascensor"
            props.creator = AUTHOR_NAME
            props.lastModifiedBy = AUTHOR_NAME
            props.description = "Dados exportados pelo SeismoLift para verificação da categoria sísmica de ascensores."
            props.keywords = "SeismoLift, EN 81-77, NP EN 1998-1, Eurocódigo 8, categoria sísmica, ascensores"
            props.category = "Cálculo sísmico de ascensores"

    def _generate_report_docx(self, path: str):
        assert self.result is not None
        r = self.result
        doc = Document()
        doc.core_properties.title = "Relatório - Categoria sísmica do elevador"
        doc.core_properties.author = AUTHOR_NAME
        doc.core_properties.subject = "Cálculo da aceleração de projecto e categoria sísmica do ascensor"
        doc.core_properties.keywords = "SeismoLift, EN 81-77, NP EN 1998-1, Eurocódigo 8, categoria sísmica, ascensores"
        doc.core_properties.comments = "Relatório automático gerado pelo SeismoLift."

        style = doc.styles["Normal"]
        style.font.name = "Courier New"
        style.font.size = Pt(10)

        if "CustomStyle" not in doc.styles:
            custom_style = doc.styles.add_style("CustomStyle", 1)
            custom_style.font.name = "Courier New"
            custom_style.font.size = Pt(16)
        if "sub_CustomStyle" not in doc.styles:
            sub_custom_style = doc.styles.add_style("sub_CustomStyle", 1)
            sub_custom_style.font.name = "Courier New"
            sub_custom_style.font.size = Pt(12)

        title = doc.add_heading("Relatório - Categoria sísmica do elevador", level=0)
        title.style = "CustomStyle"
        for run in title.runs:
            run.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        footer = doc.sections[0].footer.paragraphs[0]
        self._add_hidden_hyperlink(footer, PROGRAM_NAME, PROGRAM_URL, size=8)
        run = footer.add_run(f" - Relatório automático - Gerado em: {now}")
        self._set_run_font(run, size=8)
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("")
        sec = doc.add_heading("1. Localização e zonamento sísmico", level=1)
        sec.style = "sub_CustomStyle"
        for run in sec.runs:
            run.bold = True
        doc.add_paragraph("")
        doc.add_paragraph(f"Local: {r.zone.concelho}, {r.zone.regiao}")

        if r.zone.z1 is None:
            doc.add_paragraph(f"Zona sísmica tipo 2 (NP EN 1998-1:2010: ANEXO NA.I): {r.zone.z2}")
            doc.add_paragraph(f"Ação sísmica tipo 2 (agR): {f'{r.zone.agR2:.3f}'.replace('.', ',')} m/s²")
        elif r.zone.z2 is None:
            doc.add_paragraph(f"Zona sísmica tipo 1 (NP EN 1998-1:2010: ANEXO NA.I): {r.zone.z1}")
            doc.add_paragraph(f"Ação sísmica tipo 1 (agR): {f'{r.zone.agR1:.3f}'.replace('.', ',')} m/s²")
        else:
            doc.add_paragraph(f"Zona sísmica tipo 1 (NP EN 1998-1:2010: ANEXO NA.I): {r.zone.z1}")
            doc.add_paragraph(f"Ação sísmica tipo 1 (agR): {f'{r.zone.agR1:.3f}'.replace('.', ',')} m/s²")
            doc.add_paragraph(f"Zona sísmica tipo 2 (NP EN 1998-1:2010: ANEXO NA.I): {r.zone.z2}")
            doc.add_paragraph(f"Ação sísmica tipo 2 (agR): {f'{r.zone.agR2:.3f}'.replace('.', ',')} m/s²")

        doc.add_paragraph(f"Acção sísmica condicionante para o cálculo: {r.zone.spectro_type}")
        if len(r.cenarios) > 1:
            for c in r.cenarios:
                doc.add_paragraph(f"Verificação {c.spectro_type}: γI = {f'{c.gamma_l:.3f}'.replace('.', ',')} [-]; agR = {f'{c.agR:.3f}'.replace('.', ',')} m/s²; ad = {f'{c.ad:.3f}'.replace('.', ',')} m/s²; categoria = {c.categoria}")
        doc.add_paragraph(f"Valor de cálculo da aceleração à superfície do terreno — ag = γI·agR: {f'{r.ag:.3f}'.replace('.', ',')} [m/s²]")

        doc.add_paragraph("")
        sec = doc.add_heading("2. Parâmetros de cálculo", level=1)
        sec.style = "sub_CustomStyle"
        for run in sec.runs:
            run.bold = True
        doc.add_paragraph("")

        desc = CLASS_DESCRIPTIONS[r.classe]
        doc.add_paragraph(f"Classe e coeficientes de importância: {r.classe} - {desc}")
        doc.add_paragraph(f"Modo de cálculo: {r.calc_mode}")
        doc.add_paragraph(f"(γI): {f'{r.gamma_l:.3f}'.replace('.', ',')} [-]")
        doc.add_paragraph(f"(γa): {f'{r.gamma_a:.3f}'.replace('.', ',')} [-] — {r.gamma_a_note}")
        doc.add_paragraph(f"Valor de cálculo da aceleração à superfície do terreno (ag): {f'{r.ag:.3f}'.replace('.', ',')} [m/s²]")
        doc.add_paragraph(f"Tipo de terreno: {self.core.terrain_label(r.terreno)}")
        doc.add_paragraph(f"Coeficiente de solo (S): {f'{r.S:.3f}'.replace('.', ',')} [-]")
        doc.add_paragraph(f"Tipo de estrutura: {r.estrutura_label}")
        doc.add_paragraph(f"(Ct): {f'{r.Ct:.3f}'.replace('.', ',')} [-]")
        doc.add_paragraph(f"Altura total do edifício (H): {f'{r.H:.3f}'.replace('.', ',')} m")
        doc.add_paragraph(f"Altura do elemento considerado (z): {f'{r.z:.3f}'.replace('.', ',')} m")
        doc.add_paragraph(f"Período fundamental do elemento não estrutural (Ta): {f'{r.Ta:.3f}'.replace('.', ',')} [s]")
        doc.add_paragraph(f"Período fundamental do edifício (T1): {f'{r.T1:.3f}'.replace('.', ',')} [s]")

        doc.add_paragraph("")
        sec = doc.add_heading("3. Desenvolvimento do cálculo", level=1)
        sec.style = "sub_CustomStyle"
        for run in sec.runs:
            run.bold = True
        doc.add_paragraph("")
        if r.calc_mode == CALC_MODE_ET11:
            doc.add_paragraph("Modo ET 11/2020 - edifício de base fixa:")
            doc.add_paragraph("ad = 1,25 · γI · agR · S · γa")
        else:
            doc.add_paragraph("Modo geral EC8 / EN 81-77:")
            doc.add_paragraph("Sa = α · S · [3(1 + z/H)/(1 + (1 - Ta/T1)²) - 0,5] ≥ α · S")
            doc.add_paragraph("ad = Sa · (γa/qa) · g")
        for c in r.cenarios:
            doc.add_paragraph(
                f"{c.spectro_type}: γI = {f'{c.gamma_l:.3f}'.replace('.', ',')}; "
                f"agR = {f'{c.agR:.3f}'.replace('.', ',')} m/s²; "
                f"ag = {f'{c.ag:.3f}'.replace('.', ',')} m/s²; "
                f"S = {f'{c.S:.3f}'.replace('.', ',')}; "
                f"Sa = {f'{c.Sa:.3f}'.replace('.', ',')}; "
                f"ad = {f'{c.ad:.3f}'.replace('.', ',')} m/s²; categoria = {c.categoria}"
            )

        doc.add_paragraph("")
        sec = doc.add_heading("4. Resultados", level=1)
        sec.style = "sub_CustomStyle"
        for run in sec.runs:
            run.bold = True
        doc.add_paragraph("")

        tabela = doc.add_table(rows=2, cols=3)
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = tabela.rows[0].cells
        hdr[0].text = "Aceleração de projeto (ad) [m/s²]"
        hdr[1].text = "Categoria sísmica do elevador"
        hdr[2].text = "Nota"

        row = tabela.rows[1].cells
        if r.categoria == 0:
            row[0].text = f"{r.ad:.3f} < 1".replace(".", ",")
            row[1].text = "0"
            row[2].text = "Não são requeridas ações adicionais."
        elif r.categoria == 1:
            row[0].text = f"1 ≤ {r.ad:.3f} < 2,5".replace(".", ",")
            row[1].text = "1"
            row[2].text = "Ações corretivas de baixa expressão."
        elif r.categoria == 2:
            row[0].text = f"2,5 ≤ {r.ad:.3f} < 4".replace(".", ",")
            row[1].text = "2"
            row[2].text = "Ações corretivas de média expressão."
        else:
            row[0].text = f"{r.ad:.3f} ≥ 4".replace(".", ",")
            row[1].text = "3"
            row[2].text = "São requeridas ações corretivas importantes."

        cores_categoria = {0: "90EE90", 1: "FFFF00", 2: "FFA500", 3: "FF6B6B"}
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cores_categoria.get(r.categoria, 'FFFFFF')))
        row[1]._tc.get_or_add_tcPr().append(shading)

        for rw in tabela.rows:
            for cell in rw.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tbl_pr = tabela._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tabela._tbl.append(tbl_pr)
        tbl_borders = parse_xml(r'''
        <w:tblBorders %s>
            <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>
            <w:insideH w:val="single" w:sz="6" w:space="0" w:color="000000"/>
            <w:insideV w:val="single" w:sz="6" w:space="0" w:color="000000"/>
        </w:tblBorders>
        ''' % nsdecls('w'))
        tbl_pr.append(tbl_borders)

        doc.add_paragraph("")
        nota_lim = doc.add_paragraph("Nota técnica: o presente relatório determina a aceleração de cálculo ad e a categoria sísmica do ascensor. A verificação completa de conformidade com a EN 81-77 requer ainda a confirmação dos requisitos construtivos, funcionais e de ensaio aplicáveis ao equipamento concreto.")
        for run in nota_lim.runs:
            run.font.size = Pt(9)

        doc.save(path)

    def _generate_report_pdf(self, path: str):
        """Gera o PDF com a mesma linguagem gráfica do relatório .docx antigo.

        Nota: ao contrário da versão anterior, esta rotina não usa a tabela extensa
        de parâmetros. O objectivo é manter o PDF visualmente alinhado com o DOCX:
        fonte monoespaçada, espaçamentos largos, secções textuais e tabela final.
        """
        assert self.result is not None
        r = self.result

        # Tenta usar Courier New no Windows; se não existir, usa Courier nativo do PDF.
        mono_regular = "Courier"
        mono_bold = "Courier-Bold"
        try:
            font_candidates = [
                r"C:\Windows\Fonts\cour.ttf",
                r"C:\Windows\Fonts\courbd.ttf",
                "/usr/share/fonts/truetype/msttcorefonts/Courier_New.ttf",
                "/usr/share/fonts/truetype/msttcorefonts/Courier_New_Bold.ttf",
            ]
            regular_path = next((fp for fp in font_candidates if os.path.exists(fp) and fp.lower().endswith(("cour.ttf", "courier_new.ttf"))), None)
            bold_path = next((fp for fp in font_candidates if os.path.exists(fp) and fp.lower().endswith(("courbd.ttf", "courier_new_bold.ttf"))), None)
            if regular_path:
                pdfmetrics.registerFont(TTFont("CourierNewSL", regular_path))
                mono_regular = "CourierNewSL"
            if bold_path:
                pdfmetrics.registerFont(TTFont("CourierNewSL-Bold", bold_path))
                mono_bold = "CourierNewSL-Bold"
        except Exception:
            mono_regular = "Courier"
            mono_bold = "Courier-Bold"

        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        def fmt(value: float) -> str:
            return f"{value:.3f}".replace(".", ",")

        def safe(text) -> str:
            return escape(str(text)).replace("\n", "<br/>")

        doc = SimpleDocTemplate(
            path,
            pagesize=A4,
            rightMargin=2.5 * cm,
            leftMargin=2.5 * cm,
            topMargin=1.7 * cm,
            bottomMargin=1.9 * cm,
        )

        title_style = ParagraphStyle(
            "SL_Title",
            fontName=mono_bold,
            fontSize=16,
            leading=20,
            alignment=0,
            spaceAfter=28,
        )
        h_style = ParagraphStyle(
            "SL_Heading",
            fontName=mono_bold,
            fontSize=12,
            leading=16,
            alignment=0,
            spaceBefore=14,
            spaceAfter=22,
        )
        n_style = ParagraphStyle(
            "SL_Normal",
            fontName=mono_regular,
            fontSize=10,
            leading=16,
            alignment=0,
            spaceAfter=8,
        )
        note_style = ParagraphStyle(
            "SL_Note",
            parent=n_style,
            fontSize=9,
            leading=12,
            spaceBefore=12,
        )
        tbl_style = ParagraphStyle(
            "SL_Table",
            fontName=mono_regular,
            fontSize=9,
            leading=11,
            alignment=1,
        )
        tbl_bold_style = ParagraphStyle(
            "SL_Table_Bold",
            fontName=mono_bold,
            fontSize=9,
            leading=11,
            alignment=1,
        )

        story = []
        story.append(Paragraph("Relatório - Categoria sísmica do elevador", title_style))

        story.append(Paragraph("1. Localização e zonamento sísmico", h_style))
        story.append(Spacer(1, 18))
        story.append(Paragraph(f"Local: {safe(r.zone.concelho)}, {safe(r.zone.regiao)}", n_style))

        if r.zone.z1 is None:
            story.append(Paragraph(f"Zona sísmica tipo 2 (NP EN 1998-1:2010: ANEXO NA.I): {safe(r.zone.z2)}", n_style))
            story.append(Paragraph(f"Ação sísmica tipo 2 (agR): {fmt(r.zone.agR2)} m/s²", n_style))
        elif r.zone.z2 is None:
            story.append(Paragraph(f"Zona sísmica tipo 1 (NP EN 1998-1:2010: ANEXO NA.I): {safe(r.zone.z1)}", n_style))
            story.append(Paragraph(f"Ação sísmica tipo 1 (agR): {fmt(r.zone.agR1)} m/s²", n_style))
        else:
            story.append(Paragraph(f"Zona sísmica tipo 1 (NP EN 1998-1:2010: ANEXO NA.I): {safe(r.zone.z1)}", n_style))
            story.append(Paragraph(f"Ação sísmica tipo 1 (agR): {fmt(r.zone.agR1)} m/s²", n_style))
            story.append(Paragraph(f"Zona sísmica tipo 2 (NP EN 1998-1:2010: ANEXO NA.I): {safe(r.zone.z2)}", n_style))
            story.append(Paragraph(f"Ação sísmica tipo 2 (agR): {fmt(r.zone.agR2)} m/s²", n_style))

        story.append(Paragraph(f"Acção sísmica condicionante para o cálculo: {safe(r.zone.spectro_type)}", n_style))
        if len(r.cenarios) > 1:
            for c in r.cenarios:
                story.append(Paragraph(
                    f"Verificação {safe(c.spectro_type)}: γI = {fmt(c.gamma_l)} [-]; agR = {fmt(c.agR)} m/s²; ad = {fmt(c.ad)} m/s²;<br/>categoria = {c.categoria}",
                    n_style,
                ))
        story.append(Paragraph(
            f"Valor de cálculo da aceleração à superfície do terreno - ag = γI·agR:<br/>{fmt(r.ag)} [m/s²]",
            n_style,
        ))

        story.append(Spacer(1, 28))
        story.append(Paragraph("2. Parâmetros de cálculo", h_style))
        story.append(Spacer(1, 18))

        desc = CLASS_DESCRIPTIONS[r.classe]
        param_lines = [
            f"Classe e coeficientes de importância (NP EN 1998-1:2010: 4.2.5; NA-4.2.5(5)P. Ver Nota & 4.3.5.3): {r.classe} - {desc}",
            f"Modo de cálculo: {r.calc_mode}",
            f"     (γI): {fmt(r.gamma_l)} [-]",
            f"     (γa): {fmt(r.gamma_a)} [-] - {r.gamma_a_note}",
            f"Valor de cálculo da aceleração à superfície do terreno (NP EN 1998-1:2010: 3.2.1 (3)) (ag): {fmt(r.ag)} [m/s²]",
            f"Tipo de terreno (NP EN 1998-1:2010: Quadro 3.1): {self.core.terrain_label(r.terreno)}",
            f"Coeficiente de solo (NP EN 1998-1:2010: Quadro 3.2/3.3 EC8) (S): {fmt(r.S)} [-]",
            f"Tipo de estrutura e coeficiente de forma (NP EN 1998-1:2010: 4.3.3.2.2 (4.6)): {r.estrutura_label}.",
            f"    (Ct): {fmt(r.Ct)} [-]",
            f"Altura total do edifício (H): {fmt(r.H)} m",
            f"Altura do elemento considerado (z): {fmt(r.z)} m",
            f"Período fundamental do elemento não estrutural (EN 81-77:2018 - Annex B) (Ta): {fmt(r.Ta)} [s]",
            f"Período fundamental do edifício na direcção relevante (EN 81-77:2018 - Annex B) (T1): {fmt(r.T1)} [s]",
        ]
        for line in param_lines:
            story.append(Paragraph(safe(line), n_style))

        story.append(Spacer(1, 28))
        story.append(Paragraph("3. Desenvolvimento do cálculo", h_style))
        story.append(Spacer(1, 18))
        if r.calc_mode == CALC_MODE_ET11:
            story.append(Paragraph("Modo ET 11/2020 - edifício de base fixa:", n_style))
            story.append(Paragraph("ad = 1,25 · γI · agR · S · γa", n_style))
        else:
            story.append(Paragraph("Modo geral EC8 / EN 81-77:", n_style))
            story.append(Paragraph("Sa = α · S · [3(1 + z/H)/(1 + (1 - Ta/T1)²) - 0,5] ≥ α · S", n_style))
            story.append(Paragraph("ad = Sa · (γa/qa) · g", n_style))
        for c in r.cenarios:
            story.append(Paragraph(
                safe(f"{c.spectro_type}: γI = {fmt(c.gamma_l)}; agR = {fmt(c.agR)} m/s²; ag = {fmt(c.ag)} m/s²; S = {fmt(c.S)}; Sa = {fmt(c.Sa)}; ad = {fmt(c.ad)} m/s²; categoria = {c.categoria}"),
                n_style,
            ))

        story.append(Spacer(1, 28))
        story.append(Paragraph("4. Resultados", h_style))
        story.append(Spacer(1, 18))

        nota = {
            0: "Não são requeridas ações adicionais.",
            1: "Ações corretivas de baixa expressão.",
            2: "Ações corretivas de média expressão.",
            3: "São requeridas ações corretivas importantes.",
        }[r.categoria]
        adtxt = {
            0: f"{fmt(r.ad)} &lt; 1",
            1: f"1 ≤ {fmt(r.ad)} &lt; 2,5",
            2: f"2,5 ≤ {fmt(r.ad)} &lt; 4",
            3: f"{fmt(r.ad)} ≥ 4",
        }[r.categoria]
        cat_color = {
            0: colors.HexColor("#90EE90"),
            1: colors.HexColor("#FFFF00"),
            2: colors.HexColor("#FFA500"),
            3: colors.HexColor("#FF6B6B"),
        }[r.categoria]

        rtab = Table(
            [
                [
                    Paragraph("Aceleração de projecto<br/>(ad) [m/s²]", tbl_style),
                    Paragraph("Categoria sísmica do<br/>elevador", tbl_style),
                    Paragraph("Nota", tbl_style),
                ],
                [
                    Paragraph(adtxt, tbl_style),
                    Paragraph(str(r.categoria), tbl_bold_style),
                    Paragraph(safe(nota), tbl_style),
                ],
            ],
            colWidths=[4.6 * cm, 4.6 * cm, 4.6 * cm],
            rowHeights=[1.45 * cm, 1.65 * cm],
            hAlign="CENTER",
        )
        rtab.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.7, colors.black),
            ("BACKGROUND", (1, 1), (1, 1), cat_color),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ]))
        story.append(rtab)
        story.append(Spacer(1, 22))
        story.append(Paragraph(
            "Nota técnica: o presente relatório determina a aceleração de cálculo ad e a categoria sísmica do ascensor. A verificação completa de conformidade com a EN 81-77 requer ainda a confirmação dos requisitos construtivos, funcionais e de ensaio aplicáveis ao equipamento concreto.",
            note_style,
        ))

        def add_footer(canvas, document):
            canvas.setTitle("Relatório - Categoria sísmica do elevador")
            canvas.setAuthor(AUTHOR_NAME)
            canvas.setSubject("Cálculo da aceleração de projecto e categoria sísmica do ascensor")
            canvas.setCreator(PROGRAM_NAME)
            canvas.setKeywords("SeismoLift, EN 81-77, NP EN 1998-1, Eurocódigo 8, categoria sísmica, ascensores")
            canvas.saveState()
            canvas.setFont(mono_regular, 8)
            footer_text = f"{PROGRAM_NAME} - Relatório automático - Gerado em: {now}"
            page_width, _ = A4
            y = 1.0 * cm
            x = (page_width - canvas.stringWidth(footer_text, mono_regular, 8)) / 2.0
            canvas.drawString(x, y, footer_text)
            # Link escondido sobre o nome do programa.
            link_w = canvas.stringWidth(PROGRAM_NAME, mono_regular, 8)
            canvas.linkURL(PROGRAM_URL, (x, y - 2, x + link_w, y + 9), relative=0, thickness=0, color=None)
            canvas.restoreState()

        doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)



class SpectrumTab(BaseTab):
    def __init__(self, parent, app: "SeismoLiftGUI"):
        super().__init__(parent, app)
        self.result: Optional[SpectrumResult] = None
        self._build()

    def _build(self):
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)

        paned = AutoPanedWindow(self, orient="horizontal", sash_position=410)
        paned.grid(row=0, column=0, sticky="nsew")

        left_scroll = ScrollableFrame(paned, width=390)
        left = left_scroll.scrollable_frame
        left.columnconfigure(0, weight=1)

        right = ttk.Frame(paned, padding=(10, 12, 12, 12))
        right.columnconfigure(0, weight=1)
        right.rowconfigure(2, weight=1)

        paned.add(left_scroll, weight=0)
        paned.add(right, weight=1)

        ttk.Label(left, text="Espetros de Resposta", style="Title.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Label(
            left,
            text="EC8: Tipo 1 / Tipo 2, parâmetros TB-TC-TD e valores espectrais",
            foreground="#555",
            wraplength=340,
        ).grid(row=1, column=0, sticky="w", pady=(0, 12))

        form = ttk.LabelFrame(left, text="Entradas", padding=10)
        form.grid(row=2, column=0, sticky="nwe")
        form.columnconfigure(1, weight=1)

        self.localidade_var = tk.StringVar()
        self.classe_var = tk.StringVar(value=self.core.class_label("II"))
        self.terreno_var = tk.StringVar(value=self.core.terrain_label("B"))
        self.estrutura_var = tk.StringVar(value="Estruturas em geral")
        self.gamma_a_var = tk.StringVar(value="Ascensor corrente (γa = 1,00)")
        self.H_var = tk.StringVar(value="12,0")
        self.z_var = tk.StringVar(value="12,0")
        self.Ta_var = tk.StringVar(value="0,0")
        self.t1_mode_var = tk.StringVar(value=T1_MODE_AUTO)
        self.T1_manual_var = tk.StringVar(value="")
        self.xi_var = tk.StringVar(value="5,0")
        self.tmax_var = tk.StringVar(value="4,0")
        self.unit_var = tk.StringVar(value="m/s²")
        self.show_type1_var = tk.BooleanVar(value=True)
        self.show_type2_var = tk.BooleanVar(value=True)

        labels = [
            "Concelho", "Classe", "Terreno", "Estrutura", "γa", "H [m]", "z [m]", "Ta [s]",
            "T1", "T1 manual [s]", "ξ [%]", "T máx [s]", "Unidade Se"
        ]
        for i, txt in enumerate(labels):
            ttk.Label(form, text=txt).grid(row=i, column=0, sticky="w", padx=(0, 8), pady=4)

        ttk.Entry(form, textvariable=self.localidade_var).grid(row=0, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Combobox(form, textvariable=self.classe_var, state="readonly", values=CLASS_OPTIONS, width=34).grid(row=1, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Combobox(form, textvariable=self.terreno_var, state="readonly", values=SOIL_OPTIONS, width=34).grid(row=2, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Combobox(form, textvariable=self.estrutura_var, state="readonly", values=list(STRUCTURE_OPTIONS_SPECTRUM.keys()), width=34).grid(row=3, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Combobox(form, textvariable=self.gamma_a_var, state="readonly", values=list(GAMMA_A_OPTIONS.keys()), width=34).grid(row=4, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Entry(form, textvariable=self.H_var).grid(row=5, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Entry(form, textvariable=self.z_var).grid(row=6, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Entry(form, textvariable=self.Ta_var).grid(row=7, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Combobox(form, textvariable=self.t1_mode_var, state="readonly", values=T1_MODE_OPTIONS, width=34).grid(row=8, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Entry(form, textvariable=self.T1_manual_var).grid(row=9, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Entry(form, textvariable=self.xi_var).grid(row=10, column=1, columnspan=2, sticky="ew", pady=4)
        ttk.Entry(form, textvariable=self.tmax_var).grid(row=11, column=1, columnspan=2, sticky="ew", pady=4)
        self.unit_combo = ttk.Combobox(form, textvariable=self.unit_var, state="readonly", values=SPECTRUM_UNIT_OPTIONS, width=34)
        self.unit_combo.grid(row=12, column=1, columnspan=2, sticky="ew", pady=4)
        self.unit_combo.bind("<<ComboboxSelected>>", lambda _event: self._refresh_results_and_plot())

        show_frame = ttk.Frame(form)
        show_frame.grid(row=13, column=0, columnspan=3, sticky="ew", pady=(6, 0))
        ttk.Checkbutton(show_frame, text="Mostrar Tipo 1", variable=self.show_type1_var, command=self._plot_results_if_any).grid(row=0, column=0, sticky="w", padx=(0, 12))
        ttk.Checkbutton(show_frame, text="Mostrar Tipo 2", variable=self.show_type2_var, command=self._plot_results_if_any).grid(row=0, column=1, sticky="w")

        buttons = ttk.Frame(left)
        buttons.grid(row=3, column=0, sticky="ew", pady=(12, 0))
        buttons.columnconfigure((0, 1), weight=1)
        ttk.Button(buttons, text="Calcular", command=self.calculate).grid(row=0, column=0, sticky="ew", padx=(0, 6))
        ttk.Button(buttons, text="Limpar", command=self.clear_all).grid(row=0, column=1, sticky="ew")
        ttk.Button(buttons, text="Exportar .xlsx", command=self.export_xlsx).grid(row=1, column=0, sticky="ew", padx=(0, 6), pady=(8, 0))
        ttk.Button(buttons, text="Exportar CSV", command=self.export_csv).grid(row=1, column=1, sticky="ew", pady=(8, 0))
        ttk.Button(buttons, text="Relatório .docx", command=lambda: self.export_spectrum_report("docx")).grid(row=2, column=0, sticky="ew", padx=(0, 6), pady=(8, 0))
        ttk.Button(buttons, text="Relatório .pdf", command=lambda: self.export_spectrum_report("pdf")).grid(row=2, column=1, sticky="ew", pady=(8, 0))
        ttk.Button(buttons, text="Guardar gráfico", command=self.save_plot).grid(row=3, column=0, columnspan=2, sticky="ew", pady=(8, 0))

        notes = ttk.LabelFrame(left, text="Notas", padding=10)
        notes.grid(row=5, column=0, sticky="new", pady=(12, 0))
        ttk.Label(
            notes,
            text=(
                "• Portugal Continental: mostra Tipo 1 e Tipo 2, quando disponíveis.\n"
                "• O gráfico assinala TB, TC, TD, T1 e Ta.\n"
                "• O XLSX exporta valores Se(T) e SDe(T) para todas as acções."
            ),
            justify="left", wraplength=340,
        ).grid(row=0, column=0, sticky="w")

        header = ttk.Frame(right)
        header.grid(row=0, column=0, sticky="ew", pady=(0, 8))
        header.columnconfigure(0, weight=1)
        ttk.Label(header, text="Resultados", style="Title.TLabel").grid(row=0, column=0, sticky="w")

        summary = ttk.LabelFrame(right, text="Resumo", padding=(8, 6))
        summary.grid(row=1, column=0, sticky="ew", pady=(0, 8))
        for col in range(4):
            summary.columnconfigure(col, weight=1)
        self.summary_vars = {
            "local": tk.StringVar(value="—"),
            "acao": tk.StringVar(value="—"),
            "T1": tk.StringVar(value="—"),
            "Sa": tk.StringVar(value="—"),
            "ad": tk.StringVar(value="—"),
            "cat": tk.StringVar(value="—"),
            "se_t1": tk.StringVar(value="—"),
            "sde_t1": tk.StringVar(value="—"),
        }
        summary_items = [
            ("Local", "local"), ("Acção condicionante", "acao"), ("T1", "T1"), ("Sa", "Sa"),
            ("ad", "ad"), ("Categoria", "cat"), ("Se(T1)", "se_t1"), ("SDe(T1)", "sde_t1"),
        ]
        for idx, (title, key) in enumerate(summary_items):
            box = ttk.Frame(summary, padding=(6, 2))
            box.grid(row=idx // 4, column=idx % 4, sticky="ew")
            ttk.Label(box, text=title, style="Small.TLabel").grid(row=0, column=0, sticky="w")
            ttk.Label(box, textvariable=self.summary_vars[key], style="Summary.TLabel").grid(row=1, column=0, sticky="w")

        split = AutoPanedWindow(right, orient="vertical", sash_position=165)
        split.grid(row=2, column=0, sticky="nsew")

        top = ttk.Frame(split)
        bottom = ttk.Frame(split)
        split.add(top, weight=1)
        split.add(bottom, weight=3)

        top.columnconfigure(0, weight=1)
        top.rowconfigure(0, weight=1)
        self.tree = ttk.Treeview(top, columns=("valor",), show="tree headings", style="Result.Treeview", height=7)
        self.tree.heading("#0", text="Parâmetro")
        self.tree.heading("valor", text="Valor")
        self.tree.column("#0", width=300, minwidth=190, anchor="w", stretch=False)
        self.tree.column("valor", width=520, minwidth=220, anchor="w", stretch=True)
        self.tree.grid(row=0, column=0, sticky="nsew")
        yscroll = ttk.Scrollbar(top, orient="vertical", command=self.tree.yview)
        xscroll = ttk.Scrollbar(top, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=yscroll.set, xscrollcommand=xscroll.set)
        yscroll.grid(row=0, column=1, sticky="ns")
        xscroll.grid(row=1, column=0, sticky="ew")

        bottom.columnconfigure(0, weight=1)
        bottom.rowconfigure(0, weight=1)
        self.figure = Figure(figsize=(6.6, 3.8), dpi=100)
        self.canvas = FigureCanvasTkAgg(self.figure, master=bottom)
        self.canvas.get_tk_widget().grid(row=0, column=0, sticky="nsew")
        self._draw_placeholder()

    def calculate(self):
        try:
            t1_mode = self.t1_mode_var.get()
            t1_manual = None
            if t1_mode == T1_MODE_MANUAL:
                t1_manual = self.parse_float(self.T1_manual_var.get(), "T1 manual")
            result = self.core.compute_spectrum(
                localidade=self.localidade_var.get().strip(),
                classe=self.classe_var.get().strip(),
                terreno=self.terreno_var.get().strip(),
                estrutura=self.estrutura_var.get().strip(),
                H=self.parse_float(self.H_var.get(), "H"),
                z=self.parse_float(self.z_var.get(), "z"),
                xi=self.parse_float(self.xi_var.get(), "ξ"),
                t_max=self.parse_float(self.tmax_var.get(), "T máx"),
                gamma_a_option=self.gamma_a_var.get(),
                Ta=self.parse_float(self.Ta_var.get(), "Ta"),
                t1_mode=t1_mode,
                T1_manual=t1_manual,
            )
        except Exception as exc:
            messagebox.showerror("Erro", str(exc))
            return

        self.result = result
        self._populate_results()
        self._plot_results()

    def _populate_results(self):
        assert self.result is not None
        r = self.result
        for item in self.tree.get_children():
            self.tree.delete(item)

        data = [
            ("Concelho", f"{r.zone.concelho}, {r.zone.regiao}"),
            ("Acção sísmica condicionante", r.zone.spectro_type),
            ("Classe de importância", self.core.class_label(r.class_importance)),
            ("agR condicionante [m/s²]", self._fmt(r.zone.agR_used)),
            ("γI condicionante [-]", self._fmt(r.gamma_I)),
            ("γa [-]", self._fmt(r.gamma_a)),
            ("Critério γa", r.gamma_a_note),
            ("ag condicionante [m/s²]", self._fmt(r.ag)),
            ("α = ag/g [-]", self._fmt(r.alpha)),
            ("Tipo de terreno", self.core.terrain_label(r.terrain_type)),
            ("Parâmetros condicionantes S / TB / TC / TD", f"{self._fmt(r.S)} / {self._fmt(r.TB)} / {self._fmt(r.TC)} / {self._fmt(r.TD)}"),
            ("ξ [%] / η [-]", f"{self._fmt(r.xi)} / {self._fmt(r.eta)}"),
            ("Ct [-]", self._fmt(r.Ct)),
            ("H [m] / z [m]", f"{self._fmt(r.H)} / {self._fmt(r.z)}"),
            ("T1 [s]", f"{self._fmt(r.T1)} ({r.t1_mode})"),
            ("Ta [s]", self._fmt(r.Ta)),
            ("Sa não estrutural [-]", self._fmt(r.Sa_ns)),
            ("ad [m/s²]", self._fmt(r.ad)),
            ("Categoria sísmica", str(r.categoria)),
        ]
        for k, v in data:
            self.tree.insert("", "end", text=k, values=(v,))

        self.tree.insert("", "end", text="--- Acções avaliadas ---", values=("",))
        for c in r.cenarios:
            self.tree.insert("", "end", text=f"{c.spectro_type}: γI / agR / ad / categoria", values=(f"{self._fmt(c.gamma_l)} / {self._fmt(c.agR)} / {self._fmt(c.ad)} / {c.categoria}",))

        self.tree.insert("", "end", text="--- Valores espectrais no T1 ---", values=("",))
        for curve in r.spectra:
            Se_t1, Sde_t1 = self.core.curve_values_at(curve, r.T1, r.eta)
            Se_txt = self._fmt(Se_t1 / G if self.unit_var.get() == "g" else Se_t1)
            unit = self.unit_var.get()
            self.tree.insert("", "end", text=f"{curve.spectro_type}: Se(T1) / SDe(T1)", values=(f"{Se_txt} {unit} / {self._fmt(Sde_t1, 5)} m",))

        if r.Ta > 0:
            self.tree.insert("", "end", text="--- Valores espectrais no Ta ---", values=("",))
            for curve in r.spectra:
                Se_ta, Sde_ta = self.core.curve_values_at(curve, r.Ta, r.eta)
                Se_txt = self._fmt(Se_ta / G if self.unit_var.get() == "g" else Se_ta)
                unit = self.unit_var.get()
                self.tree.insert("", "end", text=f"{curve.spectro_type}: Se(Ta) / SDe(Ta)", values=(f"{Se_txt} {unit} / {self._fmt(Sde_ta, 5)} m",))

        self._update_summary()

    def _update_summary(self):
        if not hasattr(self, "summary_vars"):
            return
        if self.result is None:
            for var in self.summary_vars.values():
                var.set("—")
            return
        r = self.result
        unit = self.unit_var.get()
        controlling_curve = next(cur for cur in r.spectra if cur.spectro_type == r.zone.spectro_type)
        se_t1, sde_t1 = self.core.curve_values_at(controlling_curve, r.T1, r.eta)
        se_display = se_t1 / G if unit == "g" else se_t1
        self.summary_vars["local"].set(f"{r.zone.concelho}")
        self.summary_vars["acao"].set(r.zone.spectro_type)
        self.summary_vars["T1"].set(f"{self._fmt(r.T1)} s")
        self.summary_vars["Sa"].set(self._fmt(r.Sa_ns))
        self.summary_vars["ad"].set(f"{self._fmt(r.ad)} m/s²")
        self.summary_vars["cat"].set(str(r.categoria))
        self.summary_vars["se_t1"].set(f"{self._fmt(se_display)} {unit}")
        self.summary_vars["sde_t1"].set(f"{self._fmt(sde_t1, 5)} m")

    def _refresh_results_and_plot(self):
        if self.result is not None:
            self._populate_results()
            self._plot_results()

    def _plot_results_if_any(self):
        if self.result is not None:
            self._plot_results()

    def _plot_results(self):
        assert self.result is not None
        r = self.result
        self.figure.clear()
        ax = self.figure.add_subplot(111)

        unit = self.unit_var.get()
        factor = 1.0 / G if unit == "g" else 1.0
        y_label = "Se(T) [g]" if unit == "g" else "Se(T) [m/s²]"

        selected = []
        if self.show_type1_var.get():
            selected.append("Tipo 1")
        if self.show_type2_var.get():
            selected.append("Tipo 2")
        if not selected:
            selected = [r.zone.spectro_type]

        for curve in r.spectra:
            if curve.spectro_type not in selected:
                continue
            is_ctrl = curve.spectro_type == r.zone.spectro_type
            label = f"Se {curve.spectro_type}" + (" (cond.)" if is_ctrl else "")
            ax.plot(
                curve.periods,
                [v * factor for v in curve.Se],
                linewidth=2.2 if is_ctrl else 1.4,
                linestyle="-" if is_ctrl else "--",
                label=label,
            )

        # Marcadores do espetro condicionante.
        for val, name in [(r.TB, "TB"), (r.TC, "TC"), (r.TD, "TD")]:
            ax.axvline(val, linestyle=":", linewidth=0.9, label=f"{name} = {val:.2f} s")
        ax.axvline(r.T1, linestyle="-.", linewidth=1.0, label=f"T1 = {r.T1:.3f} s")
        if r.Ta > 0:
            ax.axvline(r.Ta, linestyle=(0, (5, 3)), linewidth=1.0, label=f"Ta = {r.Ta:.3f} s")

        ax.set_title(f"Espetro de resposta elástico - {r.zone.concelho}", fontsize=10)
        ax.set_xlabel("Período T [s]", fontsize=9)
        ax.set_ylabel(y_label, fontsize=9)
        ax.tick_params(labelsize=8)
        ax.grid(True, linewidth=0.5)
        ax.legend(loc="best", fontsize=8, framealpha=0.9)
        self.figure.tight_layout(pad=1.0)
        self.canvas.draw_idle()

    def _draw_placeholder(self):
        self.figure.clear()
        ax = self.figure.add_subplot(111)
        ax.text(0.5, 0.5, "Preencha os dados e clique em 'Calcular'", ha="center", va="center")
        ax.set_axis_off()
        self.figure.tight_layout()
        self.canvas.draw_idle()

    def clear_all(self):
        self.localidade_var.set("")
        self.classe_var.set(self.core.class_label("II"))
        self.terreno_var.set(self.core.terrain_label("B"))
        self.estrutura_var.set("Estruturas em geral")
        self.gamma_a_var.set("Ascensor corrente (γa = 1,00)")
        self.H_var.set("12,0")
        self.z_var.set("12,0")
        self.Ta_var.set("0,0")
        self.t1_mode_var.set(T1_MODE_AUTO)
        self.T1_manual_var.set("")
        self.xi_var.set("5,0")
        self.tmax_var.set("4,0")
        self.unit_var.set("m/s²")
        self.show_type1_var.set(True)
        self.show_type2_var.set(True)
        self.result = None
        for item in self.tree.get_children():
            self.tree.delete(item)
        self._update_summary()
        self._draw_placeholder()

    def _spectrum_dataframe(self) -> pd.DataFrame:
        assert self.result is not None
        r = self.result
        data = {"T [s]": r.periods}
        for curve in r.spectra:
            data[f"Se {curve.spectro_type} [m/s²]"] = curve.Se
            data[f"Se {curve.spectro_type} [g]"] = [v / G for v in curve.Se]
            data[f"SDe {curve.spectro_type} [m]"] = curve.Sde
        return pd.DataFrame(data)

    def export_xlsx(self):
        if not self.result:
            messagebox.showwarning("Sem dados", "Calcule primeiro o espetro.")
            return
        path = filedialog.asksaveasfilename(
            title="Guardar espetro em Excel",
            defaultextension=".xlsx",
            initialfile=f"SeismoLift_espetros_{self.result.zone.concelho.replace(' ', '_')}.xlsx",
            filetypes=[("Livro Excel", "*.xlsx")],
        )
        if not path:
            return
        try:
            self._generate_spectrum_xlsx(path)
            messagebox.showinfo("Sucesso", f"Espetro guardado em:\n{path}")
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível guardar o Excel.\n\n{exc}")

    def _generate_spectrum_xlsx(self, path: str):
        assert self.result is not None
        r = self.result
        summary = [
            ["Programa", PROGRAM_NAME],
            ["Local", f"{r.zone.concelho}, {r.zone.regiao}"],
            ["Classe de importância", self.core.class_label(r.class_importance)],
            ["Tipo de terreno", self.core.terrain_label(r.terrain_type)],
            ["Estrutura", r.structure_type],
            ["Acção condicionante", r.zone.spectro_type],
            ["γI condicionante", r.gamma_I],
            ["γa", r.gamma_a],
            ["ξ [%]", r.xi],
            ["η", r.eta],
            ["H [m]", r.H],
            ["z [m]", r.z],
            ["T1 [s]", r.T1],
            ["Modo T1", r.t1_mode],
            ["Ta [s]", r.Ta],
            ["Sa não estrutural [-]", r.Sa_ns],
            ["ad [m/s²]", r.ad],
            ["Categoria sísmica", r.categoria],
        ]
        actions = [["Acção", "Zona", "γI", "agR [m/s²]", "ag [m/s²]", "S", "TB [s]", "TC [s]", "TD [s]", "Sa", "ad [m/s²]", "Categoria"]]
        for c in r.cenarios:
            curve = next(cur for cur in r.spectra if cur.spectro_type == c.spectro_type)
            actions.append([c.spectro_type, c.z_sismica, c.gamma_l, c.agR, c.ag, curve.S, curve.TB, curve.TC, curve.TD, c.Sa, c.ad, c.categoria])

        with pd.ExcelWriter(path, engine="openpyxl") as writer:
            pd.DataFrame(summary, columns=["Campo", "Valor"]).to_excel(writer, index=False, sheet_name="Resumo")
            pd.DataFrame(actions[1:], columns=actions[0]).to_excel(writer, index=False, sheet_name="Acções")
            self._spectrum_dataframe().to_excel(writer, index=False, sheet_name="Espetros")

            props = writer.book.properties
            props.title = "SeismoLift - Espetros de resposta"
            props.subject = "Espetros de resposta EC8 e valores espectrais"
            props.creator = AUTHOR_NAME
            props.lastModifiedBy = AUTHOR_NAME
            props.description = "Espetros exportados pelo SeismoLift para análise sísmica e verificação de ascensores."
            props.keywords = "SeismoLift, EN 81-77, NP EN 1998-1, Eurocódigo 8, espetro de resposta, ascensores"
            props.category = "Espetros de resposta sísmica"

    def _plot_png_bytes(self) -> BytesIO:
        self._plot_results()
        bio = BytesIO()
        self.figure.savefig(bio, format="png", dpi=160, bbox_inches="tight")
        bio.seek(0)
        return bio

    def export_spectrum_report(self, fmt: str = "docx"):
        if not self.result:
            messagebox.showwarning("Sem dados", "Calcule primeiro o espetro.")
            return
        ext = ".docx" if fmt == "docx" else ".pdf"
        filetypes = [("Documento Word", "*.docx")] if fmt == "docx" else [("Documento PDF", "*.pdf")]
        path = filedialog.asksaveasfilename(
            title="Guardar relatório dos espetros",
            defaultextension=ext,
            initialfile=f"SeismoLift_espetros_{self.result.zone.concelho.replace(' ', '_')}{ext}",
            filetypes=filetypes,
        )
        if not path:
            return
        try:
            if fmt == "docx":
                self._generate_spectrum_docx(path)
            else:
                self._generate_spectrum_pdf(path)
            messagebox.showinfo("Sucesso", f"Relatório guardado em:\n{path}")
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível guardar o relatório.\n\n{exc}")

    def _generate_spectrum_docx(self, path: str):
        assert self.result is not None
        r = self.result
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

        cp = doc.core_properties
        cp.title = "SeismoLift - Relatório de espetros de resposta"
        cp.subject = "Espetros de resposta EC8"
        cp.author = AUTHOR_NAME
        cp.keywords = "SeismoLift, NP EN 1998-1, Eurocódigo 8, espetro de resposta"
        cp.comments = "Relatório automático gerado pelo SeismoLift."

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Relatório - Espetros de resposta")
        self._set_run_font(run, "Courier New", 13, True)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_hidden_hyperlink(p, PROGRAM_NAME, PROGRAM_URL, "Courier New", 8, False)
        run = p.add_run(f" {PROGRAM_VERSION} - Relatório automático - Gerado em: {datetime.now():%Y-%m-%d %H:%M:%S}")
        self._set_run_font(run, "Courier New", 8, False)

        for heading in ["1. Dados principais"]:
            h = doc.add_paragraph()
            self._set_run_font(h.add_run(heading), "Courier New", 10, True)

        rows = [
            ("Local", f"{r.zone.concelho}, {r.zone.regiao}"),
            ("Classe de importância", self.core.class_label(r.class_importance)),
            ("Tipo de terreno", self.core.terrain_label(r.terrain_type)),
            ("Estrutura", r.structure_type),
            ("Acção condicionante", r.zone.spectro_type),
            ("γI / γa", f"{self._fmt(r.gamma_I)} / {self._fmt(r.gamma_a)}"),
            ("H / z", f"{self._fmt(r.H)} m / {self._fmt(r.z)} m"),
            ("T1 / Ta", f"{self._fmt(r.T1)} s / {self._fmt(r.Ta)} s"),
            ("ξ / η", f"{self._fmt(r.xi)} % / {self._fmt(r.eta)}"),
            ("Sa não estrutural", self._fmt(r.Sa_ns)),
            ("ad", f"{self._fmt(r.ad)} m/s²"),
            ("Categoria sísmica", str(r.categoria)),
        ]
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Parâmetro"
        hdr[1].text = "Valor"
        for key, val in rows:
            cells = table.add_row().cells
            cells[0].text = key
            cells[1].text = val
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for para in cell.paragraphs:
                    for run in para.runs:
                        self._set_run_font(run, "Courier New", 8)

        h = doc.add_paragraph()
        self._set_run_font(h.add_run("2. Gráfico dos espetros"), "Courier New", 10, True)
        img = self._plot_png_bytes()
        doc.add_picture(img, width=Inches(6.7))

        h = doc.add_paragraph()
        self._set_run_font(h.add_run("3. Acções avaliadas"), "Courier New", 10, True)
        action_table = doc.add_table(rows=1, cols=6)
        action_table.style = "Table Grid"
        headers = ["Acção", "Zona", "γI", "agR [m/s²]", "ad [m/s²]", "Categoria"]
        for idx, text in enumerate(headers):
            action_table.rows[0].cells[idx].text = text
        for c in r.cenarios:
            row = action_table.add_row().cells
            vals = [c.spectro_type, self._fmt(c.z_sismica), self._fmt(c.gamma_l), self._fmt(c.agR), self._fmt(c.ad), str(c.categoria)]
            for idx, text in enumerate(vals):
                row[idx].text = text
        for row in action_table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        self._set_run_font(run, "Courier New", 7)

        doc.save(path)

    def _generate_spectrum_pdf(self, path: str):
        assert self.result is not None
        r = self.result
        doc = SimpleDocTemplate(
            path, pagesize=A4,
            rightMargin=1.7 * cm, leftMargin=1.7 * cm,
            topMargin=1.5 * cm, bottomMargin=1.5 * cm,
            title="SeismoLift - Relatório de espetros de resposta",
            author=AUTHOR_NAME,
            subject="Espetros de resposta EC8",
            creator=PROGRAM_NAME,
            keywords="SeismoLift, NP EN 1998-1, Eurocódigo 8, espetro de resposta",
        )
        styles = getSampleStyleSheet()
        mono = "Courier"
        title_style = ParagraphStyle("SLTitle", parent=styles["Title"], fontName=mono, fontSize=13, leading=16, alignment=1)
        h_style = ParagraphStyle("SLHeading", parent=styles["Heading2"], fontName=mono, fontSize=10, leading=13, spaceBefore=10)
        body_style = ParagraphStyle("SLBody", parent=styles["BodyText"], fontName=mono, fontSize=8, leading=10)
        story = [Paragraph("Relatório - Espetros de resposta", title_style), Spacer(1, 8)]
        story.append(Paragraph(f"{PROGRAM_NAME} {PROGRAM_VERSION} - Relatório automático - Gerado em: {datetime.now():%Y-%m-%d %H:%M:%S}", body_style))
        story.append(Spacer(1, 12))
        story.append(Paragraph("1. Dados principais", h_style))
        rows = [["Parâmetro", "Valor"],
            ["Local", f"{r.zone.concelho}, {r.zone.regiao}"],
            ["Classe", self.core.class_label(r.class_importance)],
            ["Terreno", self.core.terrain_label(r.terrain_type)],
            ["Acção condicionante", r.zone.spectro_type],
            ["γI / γa", f"{self._fmt(r.gamma_I)} / {self._fmt(r.gamma_a)}"],
            ["T1 / Ta", f"{self._fmt(r.T1)} s / {self._fmt(r.Ta)} s"],
            ["Sa / ad", f"{self._fmt(r.Sa_ns)} / {self._fmt(r.ad)} m/s²"],
            ["Categoria", str(r.categoria)],
        ]
        table = Table(rows, colWidths=[6.0 * cm, 10.2 * cm])
        table.setStyle(TableStyle([
            ("FONTNAME", (0,0), (-1,-1), mono), ("FONTSIZE", (0,0), (-1,-1), 7.5),
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey), ("GRID", (0,0), (-1,-1), 0.4, colors.black),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ]))
        story.extend([table, Spacer(1, 12), Paragraph("2. Gráfico dos espetros", h_style)])
        img = self._plot_png_bytes()
        story.append(RLImage(img, width=16.2 * cm, height=9.6 * cm))
        story.append(Spacer(1, 10))
        story.append(Paragraph("3. Acções avaliadas", h_style))
        actions = [["Acção", "Zona", "γI", "agR", "ad", "Cat."]]
        for c in r.cenarios:
            actions.append([c.spectro_type, self._fmt(c.z_sismica), self._fmt(c.gamma_l), self._fmt(c.agR), self._fmt(c.ad), str(c.categoria)])
        t2 = Table(actions, colWidths=[3.0*cm, 2.5*cm, 2.0*cm, 3.0*cm, 3.0*cm, 1.8*cm])
        t2.setStyle(TableStyle([
            ("FONTNAME", (0,0), (-1,-1), mono), ("FONTSIZE", (0,0), (-1,-1), 7.5),
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey), ("GRID", (0,0), (-1,-1), 0.4, colors.black),
            ("ALIGN", (1,1), (-1,-1), "CENTER"),
        ]))
        story.append(t2)

        def footer(canvas, _doc):
            canvas.saveState()
            canvas.setFont(mono, 7)
            text = f"{PROGRAM_NAME} {PROGRAM_VERSION} - Engº Lutonda Tomalela"
            y = 0.9 * cm
            x = A4[0] / 2 - canvas.stringWidth(text, mono, 7) / 2
            canvas.drawString(x, y, text)
            link_w = canvas.stringWidth(PROGRAM_NAME, mono, 7)
            canvas.linkURL(PROGRAM_URL, (x, y - 2, x + link_w, y + 8), relative=0, thickness=0, color=None)
            canvas.restoreState()

        doc.build(story, onFirstPage=footer, onLaterPages=footer)

    def export_csv(self):
        if not self.result:
            messagebox.showwarning("Sem dados", "Calcule primeiro o espetro.")
            return
        path = filedialog.asksaveasfilename(
            title="Guardar espetro em CSV",
            defaultextension=".csv",
            initialfile="SeismoLift_espetros.csv",
            filetypes=[("CSV", "*.csv")],
        )
        if not path:
            return
        try:
            self._spectrum_dataframe().to_csv(path, index=False, encoding="utf-8-sig")
            messagebox.showinfo("Sucesso", f"CSV guardado em:\n{path}")
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível guardar o CSV.\n\n{exc}")

    def save_plot(self):
        if not self.result:
            messagebox.showwarning("Sem dados", "Calcule primeiro o espetro.")
            return
        path = filedialog.asksaveasfilename(
            title="Guardar gráfico",
            defaultextension=".png",
            initialfile="SeismoLift_espetro.png",
            filetypes=[("PNG", "*.png"), ("PDF", "*.pdf"), ("SVG", "*.svg")],
        )
        if not path:
            return
        try:
            self.figure.savefig(path, bbox_inches="tight")
            messagebox.showinfo("Sucesso", f"Gráfico guardado em:\n{path}")
        except Exception as exc:
            messagebox.showerror("Erro", f"Não foi possível guardar o gráfico.\n\n{exc}")


class SeismoLiftGUI(tk.Tk):
    def __init__(self, core: SeismoLiftCore):
        super().__init__()
        self.core = core
        self.title(APP_TITLE)
        self._set_app_icon()
        self.geometry("1180x740")
        self.minsize(820, 560)
        self._build_style()
        self._build()


    def _set_app_icon(self):
        """Define o ícone da janela e prepara uma imagem pequena para a barra superior."""
        self._app_icon_image = None
        self._topbar_icon_image = None
        try:
            if os.path.exists(ICON_PNG_LARGE):
                self._app_icon_image = tk.PhotoImage(file=ICON_PNG_LARGE)
                self.iconphoto(True, self._app_icon_image)
            elif os.path.exists(ICON_PNG):
                self._app_icon_image = tk.PhotoImage(file=ICON_PNG)
                self.iconphoto(True, self._app_icon_image)
        except tk.TclError:
            self._app_icon_image = None
        try:
            if os.name == "nt" and os.path.exists(ICON_ICO):
                self.iconbitmap(ICON_ICO)
        except tk.TclError:
            pass
        try:
            if os.path.exists(ICON_PNG):
                self._topbar_icon_image = tk.PhotoImage(file=ICON_PNG)
        except tk.TclError:
            self._topbar_icon_image = None

    def _build_style(self):
        style = ttk.Style(self)
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass
        style.configure("Title.TLabel", font=("Segoe UI", 14, "bold"))
        style.configure("Small.TLabel", font=("Segoe UI", 8), foreground="#666666")
        style.configure("Summary.TLabel", font=("Segoe UI", 9, "bold"))
        style.configure("Result.Treeview", rowheight=22)

    def _build(self):
        self.columnconfigure(0, weight=1)
        self.rowconfigure(1, weight=1)

        topbar = ttk.Frame(self, padding=(12, 12, 12, 0))
        topbar.grid(row=0, column=0, sticky="ew")
        topbar.columnconfigure(1, weight=1)
        if self._topbar_icon_image is not None:
            ttk.Label(topbar, image=self._topbar_icon_image).grid(row=0, column=0, sticky="w", padx=(0, 8))
        ttk.Label(topbar, text=PROGRAM_NAME, style="Title.TLabel").grid(row=0, column=1, sticky="w")

        self.notebook = ttk.Notebook(self)
        self.notebook.grid(row=1, column=0, sticky="nsew", padx=8, pady=8)

        self.tab_elevator = ElevatorTab(self.notebook, self)
        self.tab_spectrum = SpectrumTab(self.notebook, self)
        self.notebook.add(self.tab_elevator, text="Categoria Sísmica de Elevadores")
        self.notebook.add(self.tab_spectrum, text="Espetros de Resposta")

    def change_excel(self):
        path = filedialog.askopenfilename(
            title="Selecionar base de dados Excel",
            filetypes=[("Excel", "*.xlsx *.xls")],
        )
        if not path:
            return
        try:
            self.core = SeismoLiftCore(path)
            self.tab_elevator.core = self.core
            self.tab_spectrum.core = self.core
            messagebox.showinfo("Sucesso", f"Base de dados actualizada:\n{path}")
        except Exception as exc:
            messagebox.showerror("Erro", str(exc))


def _load_default_or_select_database() -> Optional[SeismoLiftCore]:
    """Carrega a base sísmica embebida/instalada; permite selecção manual apenas em caso de falha."""
    try:
        return SeismoLiftCore(DEFAULT_XLSX)
    except Exception as exc:
        root = tk.Tk()
        root.withdraw()
        messagebox.showwarning(
            "Base sísmica não encontrada",
            "Não foi possível carregar automaticamente a base sísmica do SeismoLift.\n\n"
            "Seleccione manualmente o ficheiro Zonas_Sismicas_PT.xlsx ou reinstale a aplicação.\n\n"
            f"Detalhe técnico: {exc}",
            parent=root,
        )
        path = filedialog.askopenfilename(
            parent=root,
            title="Selecionar base sísmica do SeismoLift",
            filetypes=[("Excel", "*.xlsx *.xls")],
        )
        root.destroy()
        if not path:
            return None
        return SeismoLiftCore(path)


def main():
    core = _load_default_or_select_database()
    if core is None:
        return
    app = SeismoLiftGUI(core)
    app.mainloop()


if __name__ == "__main__":
    main()
